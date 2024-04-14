from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx import Document
from babel import numbers
import docx
from sympy import re
import raz


def find_paragraph(doc, target_text):
    for i, paragraph in enumerate(doc.paragraphs):
        if target_text in paragraph.text:
            p = doc.paragraphs[i]
            p.clear()  # Удаление текста из параграфа
            return i
    return -1


def replace_text(path,  # Адрес файла
                 _cal_,  # Календарь
                 _nomDog_,  # Номер договора
                 _predmetDog_,  # Предмет договора
                 _adressDog_,  # Адрес подъемника
                 _fin_,  # Финансирование
                 _cost_,  # Сумма по договору
                 sel_value,  # Тип договора # indicate
                 _unp_,  # УНП
                 _okpo_,  # ОКПО
                 _nameOrg_,  # Наименование организации
                 _adresOrg_,  # Адрес Организации
                 _nameOrgSokr_,  # Сокращенное Название Организации
                 _fioDir_,  # ФИО Директора
                 _doljnost_,  # Должность руководителя
                 _osnovanie_,  # Основание
                 _email_,  # Электронная почта
                 _phone_,  # Телефон
                 _nameBank_,  # Наименование Банка
                 _adresBank_,  # Адрес Банка
                 _iban_,  # IBAN
                 _swift_  # SWIFT
                 ):
    # -----Вносим тип договора----------------
    _text_ = raz.get_text(sel_value, _fin_, _predmetDog_, _adressDog_)
    doc = docx.Document(path)
    index = find_paragraph(doc, "[_text_]")
    for text in reversed(_text_):
        text_paragraph = doc.paragraphs[index]
        new_paragraph = text_paragraph.insert_paragraph_before(text)
        new_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    doc.save(path)  # Сохраняем документ после вставки новых параграфов
    replace_in_tables(path, _nomDog_, _nameOrg_,
                      _fioDir_, _adresOrg_, _doljnost_, _nameOrgSokr_)


# -----Вносим правки в договор-----------
def replace_in_tables(path, _nomDog_, _nameOrg_, _fioDir_, _adresOrg_, _doljnost_, _nameOrgSokr_):
    doc = docx.Document(path)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if "[_nomDog_]" in paragraph.text:
                        paragraph.text = paragraph.text.replace(
                            "[_nomDog_]", _nomDog_)
                    if "[_nameOrg_]" in paragraph.text:
                        paragraph.text = paragraph.text.replace(
                            "[_nameOrg_]", _nameOrg_)
                    if "[_fioDir_]" in paragraph.text:
                        paragraph.text = paragraph.text.replace(
                            "[_fioDir_]", _fioDir_)
                    if "[_adresOrg_]" in paragraph.text:
                        paragraph.text = paragraph.text.replace(
                            "[_adresOrg_]", _adresOrg_)
                    if "[_doljnost_]" in paragraph.text:
                        paragraph.text = paragraph.text.replace(
                            "[_doljnost_]", _doljnost_)
                    if "[_nameOrgSokr_]" in paragraph.text:
                        paragraph.text = paragraph.text.replace(
                            "[_nameOrgSokr_]", _nameOrgSokr_)
    doc.save(path)  # Сохраняем документ после замены текста
